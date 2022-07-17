# Скрипт импортирует указанную содержимое нужной ячейки из Excel в Word.
# По умолчанию сканируется вся директория рядом со скриптом на предмет xlsx файлов.
# Номер ячейки задаётся в cell.txt.

# параметры компиляции для pyinstaller
# pyinstaller export.py -F -i "giraffe.ico" --version-file info.txt

from docx import Document
from openpyxl import load_workbook
from os import walk
from progress.bar import IncrementalBar
import time


path = '.'
word_name = 'test.docx'


##### проверяем наличие файлика с названием ячейки, если нет - создаём
try:
    with open('cell.txt', 'r') as file:
        source = file.readline()
except:
    print('Создаю файл cell.txt, открой его и задай название ячейки')
    with open('cell.txt', 'w') as file:
        file.writelines(['G14'])
    time.sleep(10)
    raise SystemExit('')

def xlsx_files(path: str):
    filenames = next(walk(path), (None, None, []))[2]                    # список всех файлов рядом со скриптом
    xlsx_files = [file for file in filenames if file[-4:] == 'xlsx']     # список только xlsx
    return xlsx_files

def extractor(filelist: list) -> list:
    ''' import from excel filelist to'''
    bar = IncrementalBar('Выпей пока кофейку, Шкатаню!', max = len(files))
    to_word = []
    for file in filelist:
        bar.next()
        try:
            wb = load_workbook(file)
            ws = wb.active
            cell = ws[source].value
        except:
            cell = None
        if not cell:
            cell = ''
        # print(f'{file}  {source}    содержит {cell}')
        to_word.append((file,cell))
    bar.finish()
    exporter(to_word)

def exporter(to_word: list):
    ''' Экспортирует данные из списка кортежей в docx
        [(имя файла, данные в ячейке),(имя файла, данные в ячейке), ...]'''
        
    document = Document()
    p = document.add_paragraph(f'Содержание ячейки ')
    p.add_run(source).bold = True
    p.add_run(' во всех документах xlsx в этой папке')

    # табличка в docx
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Имя файла'
    hdr_cells[1].text = 'Содержание ячейки'
    for name, content in to_word:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = str(content)
    document.add_page_break()
    print('Export to Word format...')
    document.save(word_name)


files = xlsx_files(path)
extractor(files)


