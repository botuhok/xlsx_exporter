# Скрипт импортирует указанную содержимое нужной ячейки из Excel в Word.
# По умолчанию сканируется вся директория рядом со скриптом на предмет xlsx файлов.
# Номер ячейки задаётся в cell.txt. line = <номер строки>, column = <название столбца>

# Доп. возможности для экспорта в docx:
# document.add_heading('Heading, level 1', level=1)
# document.add_paragraph('Intense quote', style='Intense Quote')
# document.add_paragraph('first item in unordered list', style='List Bullet')
# document.add_paragraph('first item in ordered list', style='List Number')
# document.add_picture('1.png', width=Inches(1.25))    # Добавление изображения

# параметры компиляции для pyinstaller
# pyinstaller export.py --onefile -F -i "giraffe.ico" --hidden-import pyexcel_io.writers.csv_in_file --hidden-import pyexcel_io.writers.csv_in_memory --hidden-import pyexcel_io.writers.csv_sheet --hidden-import pyexcel_io.writers.csvz_sheet --hidden-import pyexcel_io.writers.csvz_writer --hidden-import pyexcel_xlsx --hidden-import pyexcel_xlsx.xlsxr --hidden-import pyexcel_xlsx.xlsxw


from docx import Document
from os import walk
from progress.bar import IncrementalBar
import pyexcel as pe
import time




##### проверяем наличие файлика с переменными, если нет - создаём
try:
    with open('cell.txt', 'r') as file:
        line = int(file.readline().split()[-1])             # строка
        column = ord(file.readline().split()[-1]) - 64      # столбец (буква -> порядковый номер)
except:
    print('Создаю файл cell.txt, открой его и задай номер строки и название столбца')
    with open('cell.txt', 'w') as file:
        file.writelines(['line = 1\n', 'column = A'])
    time.sleep(10)
    exit()

path = '.'



######## import from excel
class Extractor:
    ''' Вытаскивает инфу из нужной ячейки xlsx 
        path - путь, где лежит(лежат) xlsx файлы.
        line - строка, columns - столбец'''
        
    def __init__(self): 
        self.path = path
        self.xlsx_files = self.lst_xlsx_files()                     # список файлов xlsx

    def lst_xlsx_files(self) -> list:
        ''' Возвращает список файлов xslx в директории self.path '''
        filenames = next(walk(self.path), (None, None, []))[2]                    # список всех файлов рядом со скриптом
        xlsx_files = [file for file in filenames if file[-4:] == 'xlsx']          # список только xlsx
        return xlsx_files
    

    def set_file(self, file:str):
        ''' Задаёт файл, который нужно обрабатывать и генерирует records '''
        self.file = file
        self.records = pe.iget_records(file_name = self.file)

    def extract(self, line: int, column: int):
        ''' вытаскиваем итерируя генератор records до нужной строки '''
        for _ in range(line - 2):
            try:
                next(self.records)
            except:
                return "С этим файлом ничего не вышло"
        lst = [value for value in next(self.records).values()]
        try:
            return lst[column - 1]
        except:
            return lst

##### export to docx
def Exporter(to_word: list):
    ''' Экспортирует данные из списка кортежей в docx
        [(имя файла, данные в ячейке),(имя файла, данные в ячейке), ...]'''
    document = Document()
    # document.add_heading('Привет, Шкатаню!', 0)
    p = document.add_paragraph(f'Содержание ячейки в строкe ')

    p.add_run(str(line)).bold = True
    p.add_run(' и столбце ')
    p.add_run(str(chr(column + 64))).bold = True
    p.add_run(' во всех документах xlsx в этой папке')

    # табличка в docx
    table = document.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Имя файла'
    hdr_cells[1].text = 'Содержание ячейки'
    for name, content in to_word:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = str(content)
    document.add_page_break()
    document.save('test.docx')





f = Extractor()
to_word = []
files = f.xlsx_files
bar = IncrementalBar('Выпей пока кофейку, Шкатаню!', max = len(files))
# циклом проходимся по всем файлам внутри нужной папки
for file in f.xlsx_files:
    bar.next()
    f.set_file(file)
    cell = f.extract(line, column)
    if isinstance(cell, list):
        cell = ' '.join(str(i) for i in cell)
    # print(f'Ячейка в строке {line} и столбце {str(chr(column + 64))} содержит {cell}')
    to_word.append((file, cell))
bar.finish()
print('Export to Word format...')
Exporter(to_word)









